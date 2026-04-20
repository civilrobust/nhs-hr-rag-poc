from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NHS ICT HR Policy Assistant')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0, 94, 184)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('System Update - Expanded Policy Coverage')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 48, 135)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('Version 2.0 - March 2026')
run.font.size = Pt(14)

# Metadata
doc.add_paragraph()
doc.add_paragraph('Author: David Sewoke')
doc.add_paragraph('Date: March 18, 2026')
doc.add_paragraph('Update Type: Major Feature Enhancement')

doc.add_page_break()

# Executive Summary
doc.add_heading('Executive Summary', 1)

doc.add_paragraph('The NHS ICT HR Policy Assistant has been significantly enhanced with the addition of 10 new comprehensive policy documents, doubling the system\'s knowledge base from 10 to 20 policies.')

doc.add_paragraph()
p = doc.add_paragraph('Key Metrics:')
run = p.runs[0]
run.bold = True

doc.add_paragraph('• Policy documents: Increased from 10 to 20 (+100%)', style='List Bullet')
doc.add_paragraph('• Indexed chunks: Increased from 85 to 159 (+87%)', style='List Bullet')
doc.add_paragraph('• Coverage areas: Expanded to include equipment, on-call, training, recruitment', style='List Bullet')
doc.add_paragraph('• System performance: Maintained fast response times despite doubled dataset', style='List Bullet')

doc.add_page_break()

# What's New
doc.add_heading('What\'s New in Version 2.0', 1)

doc.add_heading('10 New Policy Documents Added', 2)

doc.add_paragraph('The following comprehensive policy documents have been added to the system:')
doc.add_paragraph()

# Policy 1
doc.add_heading('1. IT Equipment Policy (HR-ICT-011)', 3)
doc.add_paragraph('Comprehensive guidance on equipment allocation, security, and management.')
doc.add_paragraph()
doc.add_paragraph('Key topics covered:')
doc.add_paragraph('• Standard equipment allocation (desktop vs mobile/hybrid staff)', style='List Bullet')
doc.add_paragraph('• Mobile phone provisions and data allowances', style='List Bullet')
doc.add_paragraph('• Home working equipment entitlements', style='List Bullet')
doc.add_paragraph('• Security requirements and acceptable use', style='List Bullet')
doc.add_paragraph('• Equipment return procedures when leaving', style='List Bullet')
doc.add_paragraph('• Loss, damage, and replacement protocols', style='List Bullet')

doc.add_paragraph()

# Policy 2
doc.add_heading('2. On-Call Policy (HR-ICT-012)', 3)
doc.add_paragraph('Fair and transparent on-call arrangements for out-of-hours support.')
doc.add_paragraph()
doc.add_paragraph('Key topics covered:')
doc.add_paragraph('• On-call allowances: £250/week primary, £150/week secondary', style='List Bullet')
doc.add_paragraph('• Call-out payments and response time requirements', style='List Bullet')
doc.add_paragraph('• Rota management and swap procedures', style='List Bullet')
doc.add_paragraph('• Health and wellbeing provisions (fatigue management)', style='List Bullet')
doc.add_paragraph('• Incident logging and payment processing', style='List Bullet')

doc.add_paragraph()

# Policy 3
doc.add_heading('3. Training & Professional Development Policy (HR-ICT-013)', 3)
doc.add_paragraph('Investment in staff development through structured training budgets.')
doc.add_paragraph()
doc.add_paragraph('Key topics covered:')
doc.add_paragraph('• Annual training budgets by band (£500-£2,000)', style='List Bullet')
doc.add_paragraph('• Study leave entitlements (up to 5 days paid)', style='List Bullet')
doc.add_paragraph('• Degree-level sponsorship (up to 50% tuition fees)', style='List Bullet')
doc.add_paragraph('• Professional body memberships (BCS, CITP, CISSP)', style='List Bullet')
doc.add_paragraph('• E-learning subscriptions (LinkedIn Learning, Pluralsight)', style='List Bullet')
doc.add_paragraph('• Conference attendance and speaker opportunities', style='List Bullet')

doc.add_paragraph()

# Policy 4
doc.add_heading('4. Probation Period Policy (HR-ICT-014)', 3)
doc.add_paragraph('Clear expectations and assessment process for new starters.')
doc.add_paragraph()
doc.add_paragraph('Key topics covered:')
doc.add_paragraph('• Standard 6-month probation period', style='List Bullet')
doc.add_paragraph('• Formal review meetings at 1, 3, and 6 months', style='List Bullet')
doc.add_paragraph('• Assessment criteria (technical skills, behavior, values)', style='List Bullet')
doc.add_paragraph('• Extension and failure procedures', style='List Bullet')
doc.add_paragraph('• Support mechanisms (buddy system, weekly 1-to-1s)', style='List Bullet')

doc.add_paragraph()

# Policy 5
doc.add_heading('5. Working Hours & Flexitime Policy (HR-ICT-015)', 3)
doc.add_paragraph('Flexible working arrangements supporting work-life balance.')
doc.add_paragraph()
doc.add_paragraph('Key topics covered:')
doc.add_paragraph('• Standard 37.5 hours per week', style='List Bullet')
doc.add_paragraph('• Core hours: 10:00-15:00 (must be present)', style='List Bullet')
doc.add_paragraph('• Flexible hours: 07:00-10:00 and 15:00-19:00', style='List Bullet')
doc.add_paragraph('• Flexitime accrual and flexi-leave', style='List Bullet')
doc.add_paragraph('• Break entitlements and overtime procedures', style='List Bullet')

doc.add_paragraph()

# Policy 6
doc.add_heading('6. Travel & Expenses Policy (HR-ICT-016)', 3)
doc.add_paragraph('Fair reimbursement for business travel and expenses.')
doc.add_paragraph()
doc.add_paragraph('Key topics covered:')
doc.add_paragraph('• Mileage rates: 45p/mile (personal vehicle)', style='List Bullet')
doc.add_paragraph('• Accommodation maximum rates (£100-£150)', style='List Bullet')
doc.add_paragraph('• Subsistence allowances (£38/day maximum)', style='List Bullet')
doc.add_paragraph('• Public transport and taxi usage guidelines', style='List Bullet')
doc.add_paragraph('• Claims process via ESR system', style='List Bullet')

doc.add_paragraph()

# Policy 7
doc.add_heading('7. IT Security & Acceptable Use Policy (HR-ICT-017)', 3)
doc.add_paragraph('Critical security requirements and acceptable use guidelines.')
doc.add_paragraph()
doc.add_paragraph('Key topics covered:')
doc.add_paragraph('• Password requirements (12+ characters, 90-day changes)', style='List Bullet')
doc.add_paragraph('• Data protection and encryption requirements', style='List Bullet')
doc.add_paragraph('• Acceptable personal use during breaks', style='List Bullet')
doc.add_paragraph('• Security incident reporting (within 2 hours)', style='List Bullet')
doc.add_paragraph('• Remote working security (VPN, public WiFi restrictions)', style='List Bullet')

doc.add_paragraph()

# Policy 8
doc.add_heading('8. Recruitment Policy (HR-ICT-018)', 3)
doc.add_paragraph('Fair and transparent recruitment process.')
doc.add_paragraph()
doc.add_paragraph('Key topics covered:')
doc.add_paragraph('• Vacancy approval and business case requirements', style='List Bullet')
doc.add_paragraph('• Advertising (NHS Jobs, internal opportunities first)', style='List Bullet')
doc.add_paragraph('• Selection process (structured interviews, technical tests)', style='List Bullet')
doc.add_paragraph('• Pre-employment checks (DBS, references, right to work)', style='List Bullet')
doc.add_paragraph('• Onboarding timeline and preparations', style='List Bullet')

doc.add_paragraph()

# Policy 9
doc.add_heading('9. Redundancy & Restructuring Policy (HR-ICT-019)', 3)
doc.add_paragraph('Fair and supportive process during organizational changes.')
doc.add_paragraph()
doc.add_paragraph('Key topics covered:')
doc.add_paragraph('• Consultation periods (30-90 days depending on numbers)', style='List Bullet')
doc.add_paragraph('• Objective selection criteria', style='List Bullet')
doc.add_paragraph('• Redeployment opportunities and trial periods', style='List Bullet')
doc.add_paragraph('• Enhanced redundancy pay (up to 30 weeks)', style='List Bullet')
doc.add_paragraph('• Career support and coaching', style='List Bullet')

doc.add_paragraph()

# Policy 10
doc.add_heading('10. Career Progression Policy (HR-ICT-020)', 3)
doc.add_paragraph('Clear pathways for career development and progression.')
doc.add_paragraph()
doc.add_paragraph('Key topics covered:')
doc.add_paragraph('• NHS band structure (Bands 3-8 explained)', style='List Bullet')
doc.add_paragraph('• Promotion criteria and minimum time in role', style='List Bullet')
doc.add_paragraph('• Acting up opportunities and allowances', style='List Bullet')
doc.add_paragraph('• Development pathways (technical, management, hybrid, project)', style='List Bullet')
doc.add_paragraph('• Secondment opportunities (internal and external)', style='List Bullet')

doc.add_page_break()

# Technical Improvements
doc.add_heading('Technical Improvements', 1)

doc.add_heading('System Performance', 2)
doc.add_paragraph('Despite doubling the knowledge base, the system maintains excellent performance:')
doc.add_paragraph()
doc.add_paragraph('• Query response time: <3 seconds average', style='List Bullet')
doc.add_paragraph('• Indexing time: 159 chunks indexed in <2 seconds', style='List Bullet')
doc.add_paragraph('• Memory usage: Optimized ChromaDB storage', style='List Bullet')
doc.add_paragraph('• Search accuracy: Improved with larger corpus', style='List Bullet')

doc.add_heading('Code Improvements', 2)
doc.add_paragraph('• Fixed stats endpoint to dynamically count policy documents', style='List Bullet')
doc.add_paragraph('• Improved frontend to fetch live statistics from backend', style='List Bullet')
doc.add_paragraph('• Enhanced error handling in RAG pipeline', style='List Bullet')
doc.add_paragraph('• Better document counting mechanism', style='List Bullet')

doc.add_page_break()

# Coverage Comparison
doc.add_heading('Coverage Comparison: Before vs After', 1)

# Create table
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Metric'
header_cells[1].text = 'Version 1.0 (Original)'
header_cells[2].text = 'Version 2.0 (Updated)'

# Data rows
row1_cells = table.rows[1].cells
row1_cells[0].text = 'Policy Documents'
row1_cells[1].text = '10'
row1_cells[2].text = '20'

row2_cells = table.rows[2].cells
row2_cells[0].text = 'Indexed Chunks'
row2_cells[1].text = '85'
row2_cells[2].text = '159'

row3_cells = table.rows[3].cells
row3_cells[0].text = 'Coverage Areas'
row3_cells[1].text = 'Basic HR (leave, discipline, performance)'
row3_cells[2].text = 'Comprehensive (+ equipment, training, career, recruitment)'

doc.add_paragraph()

doc.add_heading('New Coverage Areas', 2)
doc.add_paragraph('The system can now answer questions about:')
doc.add_paragraph()
doc.add_paragraph('✅ IT equipment allocation and specifications', style='List Bullet')
doc.add_paragraph('✅ On-call allowances and call-out payments', style='List Bullet')
doc.add_paragraph('✅ Training budgets by band and study leave', style='List Bullet')
doc.add_paragraph('✅ Probation period expectations and reviews', style='List Bullet')
doc.add_paragraph('✅ Flexitime and working hours arrangements', style='List Bullet')
doc.add_paragraph('✅ Travel expense claims and mileage rates', style='List Bullet')
doc.add_paragraph('✅ IT security requirements and password policies', style='List Bullet')
doc.add_paragraph('✅ Recruitment process and pre-employment checks', style='List Bullet')
doc.add_paragraph('✅ Redundancy procedures and support', style='List Bullet')
doc.add_paragraph('✅ Career progression pathways and secondments', style='List Bullet')

doc.add_page_break()

# Example Questions
doc.add_heading('Example Questions the System Can Now Answer', 1)

doc.add_heading('Equipment & IT', 2)
doc.add_paragraph('• "What laptop will I get as a mobile worker?"')
doc.add_paragraph('• "Can I get a second monitor for home working?"')
doc.add_paragraph('• "What happens if I lose my work mobile phone?"')
doc.add_paragraph('• "How often does IT equipment get refreshed?"')

doc.add_heading('On-Call & Availability', 2)
doc.add_paragraph('• "What is the on-call allowance for primary on-call?"')
doc.add_paragraph('• "How quickly do I need to respond to a P1 incident?"')
doc.add_paragraph('• "What do I get paid if called out at 2am?"')
doc.add_paragraph('• "Can I swap my on-call week with a colleague?"')

doc.add_heading('Training & Development', 2)
doc.add_paragraph('• "What is my training budget as a Band 6?"')
doc.add_paragraph('• "Can I get AWS certification paid for?"')
doc.add_paragraph('• "How many days study leave am I entitled to?"')
doc.add_paragraph('• "Does the Trust fund BCS membership?"')

doc.add_heading('Career & Progression', 2)
doc.add_paragraph('• "What is the difference between Band 5 and Band 6?"')
doc.add_paragraph('• "How long do I need to be in role before promotion?"')
doc.add_paragraph('• "What is an acting up allowance?"')
doc.add_paragraph('• "Can I do a secondment to NHS Digital?"')

doc.add_heading('Recruitment & Probation', 2)
doc.add_paragraph('• "How long is the probation period?"')
doc.add_paragraph('• "What happens at my 3-month probation review?"')
doc.add_paragraph('• "What pre-employment checks are required?"')
doc.add_paragraph('• "How are interviews structured?"')

doc.add_page_break()

# Implementation Notes
doc.add_heading('Implementation Notes', 1)

doc.add_heading('How to Update the System', 2)
doc.add_paragraph('The Administrator Guide (separate document) contains full instructions, but briefly:')
doc.add_paragraph()
doc.add_paragraph('1. Add new .txt policy files to /data/policies/')
doc.add_paragraph('2. Delete ChromaDB index: rm -rf /data/chroma_db/')
doc.add_paragraph('3. Restart backend: python3 api.py')
doc.add_paragraph('4. System automatically re-indexes all documents')

doc.add_heading('Maintenance', 2)
doc.add_paragraph('• All 20 policies are now in version control (GitHub)')
doc.add_paragraph('• Backup performed: Commit a192802')
doc.add_paragraph('• No changes required to frontend or infrastructure')
doc.add_paragraph('• System remains fully local (no cloud dependencies)')

doc.add_page_break()

# Future Enhancements
doc.add_heading('Potential Future Enhancements', 1)

doc.add_paragraph('Based on the successful expansion to 20 policies, future versions could include:')
doc.add_paragraph()

doc.add_heading('Additional Policy Areas', 2)
doc.add_paragraph('• Whistleblowing policy', style='List Bullet')
doc.add_paragraph('• Occupational health procedures', style='List Bullet')
doc.add_paragraph('• Union representation rights', style='List Bullet')
doc.add_paragraph('• Special leave (jury duty, bereavement, etc.)', style='List Bullet')
doc.add_paragraph('• Pension and retirement guidance', style='List Bullet')

doc.add_heading('Technical Enhancements', 2)
doc.add_paragraph('• Multi-turn conversations (follow-up questions)', style='List Bullet')
doc.add_paragraph('• Policy comparison feature ("compare sick leave vs annual leave")', style='List Bullet')
doc.add_paragraph('• Export answers to PDF for record keeping', style='List Bullet')
doc.add_paragraph('• User feedback system (thumbs up/down on answers)', style='List Bullet')
doc.add_paragraph('• Analytics dashboard (most asked questions)', style='List Bullet')

doc.add_page_break()

# Cost Analysis Update
doc.add_heading('Cost Analysis - Avoiding Cloud Hosting', 1)

doc.add_paragraph('This expansion reinforces the value of the local deployment approach:')
doc.add_paragraph()

doc.add_heading('Cost Avoidance', 2)
p = doc.add_paragraph()
run = p.add_run('Linode VPS option (considered and rejected):')
run.bold = True

doc.add_paragraph('• Monthly cost: £28')
doc.add_paragraph('• Annual cost: £336')
doc.add_paragraph('• Storage: 50GB (would be 80% full after model downloads)')
doc.add_paragraph('• RAM: 2GB (insufficient for smooth operation)')
doc.add_paragraph('• Performance: Slow query responses')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Local deployment (chosen approach):')
run.bold = True
run.font.color.rgb = RGBColor(0, 170, 0)

doc.add_paragraph('• Monthly cost: £0')
doc.add_paragraph('• Annual cost: £0')
doc.add_paragraph('• Storage: Abundant (TB available)')
doc.add_paragraph('• RAM: 64GB (more than sufficient)')
doc.add_paragraph('• GPU: RTX 5080 16GB (excellent performance)')
doc.add_paragraph('• Performance: Sub-3-second query responses')

doc.add_paragraph()
p = doc.add_paragraph('Total savings: £336/year')
run = p.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 170, 0)

doc.add_page_break()

# Conclusion
doc.add_heading('Conclusion', 1)

doc.add_paragraph('Version 2.0 represents a major enhancement to the NHS ICT HR Policy Assistant, doubling the knowledge base while maintaining excellent performance and zero ongoing costs.')

doc.add_paragraph()
doc.add_paragraph('The system now provides comprehensive coverage across:')
doc.add_paragraph('• Employee lifecycle (recruitment → career progression → leaving)', style='List Bullet')
doc.add_paragraph('• Day-to-day operations (equipment, IT security, working hours)', style='List Bullet')
doc.add_paragraph('• Development and support (training, on-call, expenses)', style='List Bullet')
doc.add_paragraph('• Employment processes (probation, discipline, grievance)', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('The POC is ready for demonstration to CogStack stakeholders.')
run = p.runs[0]
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph('--- End of Update Document ---')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.runs[0]
run.italic = True

doc.save('NHS_HR_RAG_Version_2.0_Update_Document.docx')
print('✅ Update document created!')
print('📄 File: NHS_HR_RAG_Version_2.0_Update_Document.docx')

