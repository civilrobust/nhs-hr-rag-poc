from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NHS ICT HR Policy Assistant')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0, 94, 184)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Realistic Deployment & Scaling Plan')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 48, 135)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('Kings College Hospital NHS Foundation Trust')
run.font.size = Pt(14)

subtitle3 = doc.add_paragraph()
subtitle3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle3.add_run('ICT Department - 200 Staff Target Audience')
run.font.size = Pt(12)

# Metadata
doc.add_paragraph()
doc.add_paragraph('Author: David Sewoke')
doc.add_paragraph('Date: March 2026')
doc.add_paragraph('Status: Proof of Concept → Production Planning')

doc.add_page_break()

# Executive Summary
doc.add_heading('Executive Summary', 1)

doc.add_paragraph('This document outlines a realistic, cost-effective deployment strategy for the NHS ICT HR Policy Assistant, based on actual usage patterns and proven POC functionality.')

doc.add_paragraph()
p = doc.add_paragraph('Key Reality Check:')
run = p.runs[0]
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph('"It\'s an HR policy app - nobody really uses these things much."')
doc.add_paragraph()

p = doc.add_paragraph('Expected Reality:')
run = p.runs[0]
run.bold = True

doc.add_paragraph('• Target audience: 200 ICT staff', style='List Bullet')
doc.add_paragraph('• Realistic concurrent users: 2-5 (not 20!)', style='List Bullet')
doc.add_paragraph('• Peak usage: New starter onboarding, annual leave queries, problem escalations', style='List Bullet')
doc.add_paragraph('• Most common questions: "How much annual leave?", "What\'s the sick pay?", "Can I work from home?"', style='List Bullet')
doc.add_paragraph('• Actual daily queries: 10-30 total (not hundreds)', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Translation: We don\'t need a supercomputer. We need something reliable, cheap, and "good enough".')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 170, 0)

doc.add_page_break()

# Current POC Status
doc.add_heading('What We\'ve Already Built (POC)', 1)

doc.add_heading('Fully Functional System Running Today', 2)

doc.add_paragraph('Hardware: Consumer-grade PC')
doc.add_paragraph('• CPU: AMD Ryzen 9 9950X (16 cores)', style='List Bullet')
doc.add_paragraph('• GPU: NVIDIA RTX 5080 16GB', style='List Bullet')
doc.add_paragraph('• RAM: 64GB DDR5', style='List Bullet')
doc.add_paragraph('• Storage: NVMe SSD', style='List Bullet')
doc.add_paragraph('• OS: Windows 11 + WSL2 Ubuntu 24.04', style='List Bullet')
doc.add_paragraph('• Cost: £0 (already owned)', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Software Stack:')
doc.add_paragraph('• LLM: Llama 3.1 8B (local, via Ollama)', style='List Bullet')
doc.add_paragraph('• Vector DB: ChromaDB (persistent storage)', style='List Bullet')
doc.add_paragraph('• Backend: FastAPI (Python)', style='List Bullet')
doc.add_paragraph('• Frontend: React + Vite', style='List Bullet')
doc.add_paragraph('• Knowledge Base: 20 HR policy documents, 159 indexed chunks', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Performance Metrics (Proven):')
doc.add_paragraph('• Query response time: 2-3 seconds', style='List Bullet')
doc.add_paragraph('• Concurrent users tested: 5 (comfortable)', style='List Bullet')
doc.add_paragraph('• Accuracy: High (citations from correct policies)', style='List Bullet')
doc.add_paragraph('• Uptime: 100% during testing', style='List Bullet')
doc.add_paragraph('• Cost per query: £0.00', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Status: WORKS. Right now. Today. On a desktop PC.')
run = p.runs[0]
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 170, 0)

doc.add_page_break()

# Realistic Usage Patterns
doc.add_heading('Realistic Usage Patterns for HR Policy App', 1)

doc.add_heading('Let\'s Be Honest About Usage', 2)

doc.add_paragraph('HR policy apps typically have:')
doc.add_paragraph()

doc.add_paragraph('Low Daily Active Users:')
doc.add_paragraph('• 200 staff total', style='List Bullet')
doc.add_paragraph('• 5-10 users per day (2.5-5% daily active rate)', style='List Bullet')
doc.add_paragraph('• 1-2 users at any one time (peak: 3-5 during onboarding weeks)', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Usage Spikes:')
doc.add_paragraph('• New starter weeks: 10-15 queries/day', style='List Bullet')
doc.add_paragraph('• January (post-Christmas leave questions): 20-30 queries/day', style='List Bullet')
doc.add_paragraph('• Appraisal season (March-April): 15-20 queries/day', style='List Bullet')
doc.add_paragraph('• Normal weeks: 5-10 queries/day', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Most Common Questions:')
doc.add_paragraph('• "How much annual leave do I get?" (asked weekly)', style='List Bullet')
doc.add_paragraph('• "What is the sick pay policy?" (asked monthly)', style='List Bullet')
doc.add_paragraph('• "Can I work from home full time?" (asked frequently)', style='List Bullet')
doc.add_paragraph('• "What is the on-call allowance?" (new staff)', style='List Bullet')
doc.add_paragraph('• "How do I claim expenses?" (when people travel)', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Translation: This is not a high-traffic system. Nobody is hammering the HR policy bot 24/7.')
run = p.runs[0]
run.italic = True

doc.add_paragraph()
p = doc.add_paragraph('Requirement: Handle 5 concurrent users comfortably, with headroom for 10-15 during onboarding.')
run = p.runs[0]
run.bold = True

doc.add_page_break()

# Deployment Options
doc.add_heading('Deployment Options: From POC to Production', 1)

doc.add_heading('Option 1: "The Repurpose" (Recommended)', 2)
p = doc.add_paragraph('Use existing decommissioned workstation from IT asset disposal')
run = p.runs[0]
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Hardware Spec (repurpose old high-spec workstation):')
doc.add_paragraph('• Dell Precision 7920 or HP Z8 G4 (2-3 years old)', style='List Bullet')
doc.add_paragraph('• CPU: Intel Xeon or AMD Threadripper (16+ cores)', style='List Bullet')
doc.add_paragraph('• GPU: NVIDIA RTX 4060 Ti 16GB or RTX 3090 24GB (from asset pool)', style='List Bullet')
doc.add_paragraph('• RAM: 64GB minimum', style='List Bullet')
doc.add_paragraph('• Storage: 1TB NVMe SSD', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Setup:')
doc.add_paragraph('• Install Ubuntu Server 24.04', style='List Bullet')
doc.add_paragraph('• Deploy POC code (already working!)', style='List Bullet')
doc.add_paragraph('• Add nginx reverse proxy', style='List Bullet')
doc.add_paragraph('• Add Let\'s Encrypt SSL certificate', style='List Bullet')
doc.add_paragraph('• Set up systemd services (auto-start on boot)', style='List Bullet')
doc.add_paragraph('• Daily automated backups', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Cost:')
p = doc.add_paragraph('• Hardware: £0 (repurposed from IT asset disposal)')
run = p.runs[0]
run.font.color.rgb = RGBColor(0, 170, 0)
doc.add_paragraph('• Software: £0 (all open source)', style='List Bullet')
doc.add_paragraph('• Setup time: 1-2 days (David + 1 senior engineer)', style='List Bullet')
doc.add_paragraph('• Electricity: ~£50/year', style='List Bullet')
doc.add_paragraph('• Maintenance: 2 hours/month (monitoring, updates)', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Total Year 1 Cost: £50')
run = p.runs[0]
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 170, 0)

doc.add_paragraph()
doc.add_paragraph('Capacity:')
doc.add_paragraph('• Concurrent users: 10-15 comfortably', style='List Bullet')
doc.add_paragraph('• Daily queries: 100+ (way more than needed)', style='List Bullet')
doc.add_paragraph('• Response time: 2-4 seconds', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Pros:')
run = p.runs[0]
run.bold = True
doc.add_paragraph('✅ Essentially free (repurposed hardware)', style='List Bullet')
doc.add_paragraph('✅ More than adequate for 200 users', style='List Bullet')
doc.add_paragraph('✅ Quick to deploy (code already works)', style='List Bullet')
doc.add_paragraph('✅ Easy to maintain', style='List Bullet')
doc.add_paragraph('✅ Can always upgrade later if needed', style='List Bullet')

p = doc.add_paragraph('Cons:')
run = p.runs[0]
run.bold = True
doc.add_paragraph('⚠️ Single point of failure (but can rebuild in hours)', style='List Bullet')
doc.add_paragraph('⚠️ Not "enterprise grade" (but who cares for an HR bot?)', style='List Bullet')

doc.add_page_break()

doc.add_heading('Option 2: "The New Buy" (If Budget Available)', 2)
p = doc.add_paragraph('Purchase dedicated server (if repurpose not possible)')
run = p.runs[0]
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Hardware Spec:')
doc.add_paragraph('• Dell PowerEdge T550 or HPE ProLiant ML350 Gen11', style='List Bullet')
doc.add_paragraph('• CPU: Intel Xeon Silver 4410Y (24 cores)', style='List Bullet')
doc.add_paragraph('• GPU: NVIDIA RTX 4000 Ada 20GB', style='List Bullet')
doc.add_paragraph('• RAM: 128GB DDR5 ECC', style='List Bullet')
doc.add_paragraph('• Storage: 2TB NVMe SSD', style='List Bullet')
doc.add_paragraph('• Redundant PSU: Yes', style='List Bullet')
doc.add_paragraph('• 3-year warranty: Included', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('UK Pricing (March 2026):')
doc.add_paragraph('• Server: £8,000', style='List Bullet')
doc.add_paragraph('• GPU: £2,500', style='List Bullet')
doc.add_paragraph('• Setup & config: £500 (internal time)', style='List Bullet')
doc.add_paragraph('• Rack mounting: £200', style='List Bullet')
doc.add_paragraph('• Total: £11,200', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Running Costs:')
doc.add_paragraph('• Electricity: ~£150/year', style='List Bullet')
doc.add_paragraph('• Maintenance: 3 hours/month monitoring', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Total Year 1 Cost: £11,350')
run = p.runs[0]
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph('Capacity:')
doc.add_paragraph('• Concurrent users: 25-30', style='List Bullet')
doc.add_paragraph('• Daily queries: 500+ (massive overkill)', style='List Bullet')
doc.add_paragraph('• Response time: <2 seconds', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Pros:')
run = p.runs[0]
run.bold = True
doc.add_paragraph('✅ Enterprise-grade hardware', style='List Bullet')
doc.add_paragraph('✅ 3-year warranty and support', style='List Bullet')
doc.add_paragraph('✅ Redundant PSU (more reliable)', style='List Bullet')
doc.add_paragraph('✅ Room to grow (could support 500+ users)', style='List Bullet')
doc.add_paragraph('✅ ECC RAM (better stability)', style='List Bullet')

p = doc.add_paragraph('Cons:')
run = p.runs[0]
run.bold = True
doc.add_paragraph('⚠️ Costs £11k (vs £0 for repurpose)', style='List Bullet')
doc.add_paragraph('⚠️ Still single point of failure', style='List Bullet')
doc.add_paragraph('⚠️ Overkill for current needs', style='List Bullet')

doc.add_page_break()

doc.add_heading('Option 3: "The Belt & Braces" (High Availability)', 2)
p = doc.add_paragraph('Two servers for redundancy (if critical service)')
run = p.runs[0]
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Only necessary if:')
doc.add_paragraph('• Service becomes business-critical (unlikely for HR policy bot)', style='List Bullet')
doc.add_paragraph('• User base expands beyond ICT to whole Trust (future)', style='List Bullet')
doc.add_paragraph('• Budget available for "nice to have"', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Setup:')
doc.add_paragraph('• 2x servers (Option 2 spec)', style='List Bullet')
doc.add_paragraph('• Active-passive failover OR active-active load balancing', style='List Bullet')
doc.add_paragraph('• Replicated ChromaDB', style='List Bullet')
doc.add_paragraph('• HAProxy load balancer', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Total Year 1 Cost: £23,000')
run = p.runs[0]
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph()
p = doc.add_paragraph('Reality Check:')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(255, 102, 0)

doc.add_paragraph('This is overkill for an HR policy app used by 200 staff. Even if one server dies, you can rebuild from backup in 2 hours. Is 2 hours downtime worth £12k extra? Probably not.')

doc.add_paragraph()
p = doc.add_paragraph('Recommendation: Start with Option 1 or 2. Add HA later if usage proves it\'s needed.')
run = p.runs[0]
run.italic = True

doc.add_page_break()

# Recommended Approach
doc.add_heading('Recommended Deployment Strategy', 1)

doc.add_heading('Phase 1: Pilot (3 months) - Prove It Works', 2)

doc.add_paragraph('Deployment:')
doc.add_paragraph('• Option 1: Repurposed workstation (£0)', style='List Bullet')
doc.add_paragraph('• Internal URL: https://hr-assistant.ict.kch.nhs.uk', style='List Bullet')
doc.add_paragraph('• Access: ICT staff only (200 people)', style='List Bullet')
doc.add_paragraph('• Announce: Email + Teams message + brown bag demo', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Success Metrics:')
doc.add_paragraph('• 20+ unique users in first month', style='List Bullet')
doc.add_paragraph('• 50+ queries total', style='List Bullet')
doc.add_paragraph('• 80%+ positive feedback ("Was this helpful?")', style='List Bullet')
doc.add_paragraph('• <5 second average response time', style='List Bullet')
doc.add_paragraph('• Zero security incidents', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('What We\'ll Learn:')
doc.add_paragraph('• Do people actually use it? (reality check)', style='List Bullet')
doc.add_paragraph('• What questions are asked most?', style='List Bullet')
doc.add_paragraph('• Which policies need better coverage?', style='List Bullet')
doc.add_paragraph('• Is Llama 3.1 8B accurate enough? (probably yes)', style='List Bullet')
doc.add_paragraph('• Any technical issues?', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Cost: £50/year + 5 days David\'s time')
run = p.runs[0]
run.bold = True

doc.add_page_break()

doc.add_heading('Phase 2: Production (6-12 months) - Make It Official', 2)

doc.add_paragraph('If Pilot Successful:')
doc.add_paragraph('• Keep same hardware (if working well)', style='List Bullet')
doc.add_paragraph('• OR upgrade to Option 2 (new server) if budget approved', style='List Bullet')
doc.add_paragraph('• Add monitoring (Prometheus + Grafana)', style='List Bullet')
doc.add_paragraph('• Add authentication (Active Directory)', style='List Bullet')
doc.add_paragraph('• Add backup automation', style='List Bullet')
doc.add_paragraph('• Document support procedures', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Expansion Considerations:')
doc.add_paragraph('• Add more policy documents (currently 20, could expand to 50+)', style='List Bullet')
doc.add_paragraph('• Support .docx and .pdf files (not just .txt)', style='List Bullet')
doc.add_paragraph('• Add analytics dashboard (usage stats)', style='List Bullet')
doc.add_paragraph('• Consider expanding to other departments', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Cost: £50-150/year (Option 1) OR £11,350 first year + £200/year (Option 2)')
run = p.runs[0]
run.bold = True

doc.add_page_break()

# Cost Comparison
doc.add_heading('Cost Comparison: DIY vs Commercial', 1)

doc.add_heading('What We Built (In-House)', 2)

# Create table
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Item'
header_cells[1].text = 'Option 1 (Repurpose)'
header_cells[2].text = 'Option 2 (New Server)'

# Data rows
row1_cells = table.rows[1].cells
row1_cells[0].text = 'Hardware'
row1_cells[1].text = '£0 (existing)'
row1_cells[2].text = '£11,200'

row2_cells = table.rows[2].cells
row2_cells[0].text = 'Software'
row2_cells[1].text = '£0 (open source)'
row2_cells[2].text = '£0 (open source)'

row3_cells = table.rows[3].cells
row3_cells[0].text = 'Setup (5 days)'
row3_cells[1].text = 'Internal time'
row3_cells[2].text = 'Internal time'

row4_cells = table.rows[4].cells
row4_cells[0].text = 'Year 1 Total'
row4_cells[1].text = '£50'
row4_cells[2].text = '£11,350'

row5_cells = table.rows[5].cells
row5_cells[0].text = '3-Year Total'
row5_cells[1].text = '£150'
row5_cells[2].text = '£11,800'

doc.add_paragraph()

doc.add_heading('Commercial Alternative (Typical HR Chatbot SaaS)', 2)

doc.add_paragraph('Pricing (UK market):')
doc.add_paragraph('• Per-user licensing: £3-8 per user per month', style='List Bullet')
doc.add_paragraph('• 200 users × £5/month average = £1,000/month', style='List Bullet')
doc.add_paragraph('• Annual cost: £12,000/year', style='List Bullet')
doc.add_paragraph('• 3-year cost: £36,000', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('What You Get:')
doc.add_paragraph('✅ Hosted solution (no hardware needed)', style='List Bullet')
doc.add_paragraph('✅ Professional support', style='List Bullet')
doc.add_paragraph('✅ Regular updates', style='List Bullet')
doc.add_paragraph('⚠️ Data leaves NHS premises (cloud-hosted)', style='List Bullet')
doc.add_paragraph('⚠️ Generic HR knowledge (not NHS-specific)', style='List Bullet')
doc.add_paragraph('⚠️ Ongoing costs forever', style='List Bullet')
doc.add_paragraph('⚠️ Limited customization', style='List Bullet')

doc.add_paragraph()

# Savings table
doc.add_heading('Savings Analysis', 2)

table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Light Grid Accent 1'

header2_cells = table2.rows[0].cells
header2_cells[0].text = 'Timeframe'
header2_cells[1].text = 'Option 1 (DIY)'
header2_cells[2].text = 'Option 2 (DIY)'
header2_cells[3].text = 'Commercial SaaS'

row1_cells = table2.rows[1].cells
row1_cells[0].text = 'Year 1'
row1_cells[1].text = '£50'
row1_cells[2].text = '£11,350'
row1_cells[3].text = '£12,000'

row2_cells = table2.rows[2].cells
row2_cells[0].text = '3 Years'
row2_cells[1].text = '£150'
row2_cells[2].text = '£11,800'
row2_cells[3].text = '£36,000'

row3_cells = table2.rows[3].cells
row3_cells[0].text = '5 Years'
row3_cells[1].text = '£250'
row3_cells[2].text = '£12,200'
row3_cells[3].text = '£60,000'

doc.add_paragraph()
p = doc.add_paragraph('3-Year Savings:')
run = p.runs[0]
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph('• Option 1 vs Commercial: Save £35,850', style='List Bullet')
doc.add_paragraph('• Option 2 vs Commercial: Save £24,200', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Even buying a brand new £11k server saves £24k over 3 years!')
run = p.runs[0]
run.font.color.rgb = RGBColor(0, 170, 0)
run.font.size = Pt(12)

doc.add_page_break()

# Technical Improvements Needed
doc.add_heading('Production Readiness Gaps', 1)

doc.add_heading('What Works Today (POC)', 2)
doc.add_paragraph('✅ Core RAG functionality', style='List Bullet')
doc.add_paragraph('✅ 20 HR policies indexed', style='List Bullet')
doc.add_paragraph('✅ Fast query responses (2-3 seconds)', style='List Bullet')
doc.add_paragraph('✅ Accurate answers with citations', style='List Bullet')
doc.add_paragraph('✅ Clean web interface', style='List Bullet')
doc.add_paragraph('✅ Conversation history', style='List Bullet')
doc.add_paragraph('✅ "New Chat" button', style='List Bullet')

doc.add_heading('What Needs Adding for Production (Priority Order)', 2)

doc.add_paragraph('High Priority (Must Have):')
doc.add_paragraph('1. Authentication (Active Directory login)', style='List Bullet')
doc.add_paragraph('2. HTTPS/SSL certificate', style='List Bullet')
doc.add_paragraph('3. Automated backups (daily)', style='List Bullet')
doc.add_paragraph('4. Basic monitoring (is it running?)', style='List Bullet')
doc.add_paragraph('5. Error logging', style='List Bullet')
doc.add_paragraph('6. Service restart on failure (systemd)', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Medium Priority (Should Have):')
doc.add_paragraph('7. Support .docx and .pdf files (not just .txt)', style='List Bullet')
doc.add_paragraph('8. Feedback mechanism (thumbs up/down)', style='List Bullet')
doc.add_paragraph('9. Usage analytics (how many queries/day)', style='List Bullet')
doc.add_paragraph('10. Admin panel (add/remove policies)', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Low Priority (Nice to Have):')
doc.add_paragraph('11. Export answer to PDF', style='List Bullet')
doc.add_paragraph('12. Email answers to self', style='List Bullet')
doc.add_paragraph('13. Policy version tracking', style='List Bullet')
doc.add_paragraph('14. Integration with ServiceNow', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Estimated Development Time:')
run = p.runs[0]
run.bold = True

doc.add_paragraph('• High priority items: 3-5 days', style='List Bullet')
doc.add_paragraph('• Medium priority items: 5-7 days', style='List Bullet')
doc.add_paragraph('• Low priority items: Optional, add later if needed', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Total: 2 weeks to make it production-ready')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 170, 0)

doc.add_page_break()

# Risks and Mitigations
doc.add_heading('Risks & Mitigations', 1)

# Risk 1
doc.add_heading('Risk 1: Nobody Uses It', 2)
doc.add_paragraph('Likelihood: Medium')
doc.add_paragraph('Impact: Low (it\'s only £50/year)')
doc.add_paragraph()
doc.add_paragraph('Mitigation:')
doc.add_paragraph('• Brown bag lunch demo (show how easy it is)', style='List Bullet')
doc.add_paragraph('• Email campaign with example questions', style='List Bullet')
doc.add_paragraph('• Add link to Teams channel', style='List Bullet')
doc.add_paragraph('• During onboarding, point new starters to it', style='List Bullet')
doc.add_paragraph('• Promote when policy changes happen', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Worst case: We spent £50 and 2 weeks. Lesson learned.')
run = p.runs[0]
run.italic = True

# Risk 2
doc.add_heading('Risk 2: Server Hardware Failure', 2)
doc.add_paragraph('Likelihood: Low')
doc.add_paragraph('Impact: Medium (service down until rebuilt)')
doc.add_paragraph()
doc.add_paragraph('Mitigation:')
doc.add_paragraph('• Daily automated backups', style='List Bullet')
doc.add_paragraph('• Documented rebuild procedure (1-2 hours)', style='List Bullet')
doc.add_paragraph('• Keep spare GPU in asset pool', style='List Bullet')
doc.add_paragraph('• Fallback: Run on David\'s PC temporarily', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Recovery Time Objective: 2 hours')
run = p.runs[0]
run.italic = True

# Risk 3
doc.add_heading('Risk 3: AI Gives Wrong Answer', 2)
doc.add_paragraph('Likelihood: Low (tested extensively)')
doc.add_paragraph('Impact: Medium (user gets incorrect HR info)')
doc.add_paragraph()
doc.add_paragraph('Mitigation:')
doc.add_paragraph('• Always show source citations', style='List Bullet')
doc.add_paragraph('• Disclaimer on every page: "AI-generated, verify with HR if critical"', style='List Bullet')
doc.add_paragraph('• Feedback button: "Was this helpful?"', style='List Bullet')
doc.add_paragraph('• Monthly review of negative feedback', style='List Bullet')
doc.add_paragraph('• Continuous testing with known questions', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Note: System cites exact policy sections. User can verify themselves.')
run = p.runs[0]
run.italic = True

# Risk 4
doc.add_heading('Risk 4: Information Governance Concerns', 2)
doc.add_paragraph('Likelihood: Low')
doc.add_paragraph('Impact: High (project blocked)')
doc.add_paragraph()
doc.add_paragraph('Mitigation:')
doc.add_paragraph('• All data on-premise (NHS servers)', style='List Bullet')
doc.add_paragraph('• No patient data involved', style='List Bullet')
doc.add_paragraph('• No personal data stored (queries not logged with usernames)', style='List Bullet')
doc.add_paragraph('• Public HR policies only (no confidential data)', style='List Bullet')
doc.add_paragraph('• DPIA completed before production', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Pre-emptive IG engagement recommended.')
run = p.runs[0]
run.italic = True

doc.add_page_break()

# Timeline
doc.add_heading('Implementation Timeline', 1)

doc.add_heading('Week 1-2: Production Prep', 2)
doc.add_paragraph('• Secure repurposed hardware (or order new server)', style='List Bullet')
doc.add_paragraph('• Install Ubuntu Server 24.04', style='List Bullet')
doc.add_paragraph('• Deploy POC code', style='List Bullet')
doc.add_paragraph('• Add authentication (Active Directory)', style='List Bullet')
doc.add_paragraph('• Set up HTTPS with Let\'s Encrypt', style='List Bullet')
doc.add_paragraph('• Configure automated backups', style='List Bullet')
doc.add_paragraph('• Add basic monitoring', style='List Bullet')

doc.add_heading('Week 3: Testing & Documentation', 2)
doc.add_paragraph('• Internal testing (5-10 ICT staff)', style='List Bullet')
doc.add_paragraph('• Write user guide', style='List Bullet')
doc.add_paragraph('• Create support documentation', style='List Bullet')
doc.add_paragraph('• IG review and approval', style='List Bullet')

doc.add_heading('Week 4: Launch', 2)
doc.add_paragraph('• Announce to ICT department', style='List Bullet')
doc.add_paragraph('• Brown bag demo lunch session', style='List Bullet')
doc.add_paragraph('• Teams message with quick start guide', style='List Bullet')
doc.add_paragraph('• Monitor usage and feedback', style='List Bullet')

doc.add_heading('Month 2-3: Optimize', 2)
doc.add_paragraph('• Review usage analytics', style='List Bullet')
doc.add_paragraph('• Add most-requested features', style='List Bullet')
doc.add_paragraph('• Add missing policies based on feedback', style='List Bullet')
doc.add_paragraph('• Refine answers based on thumbs down feedback', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Go-Live Target: 4 weeks from approval')
run = p.runs[0]
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 94, 184)

doc.add_page_break()

# Conclusion
doc.add_heading('Conclusion: Keep It Simple', 1)

doc.add_paragraph('We\'ve built a working AI-powered HR policy assistant in a weekend. It runs on a desktop PC. It costs nothing to operate. It answers questions accurately in 2-3 seconds.')

doc.add_paragraph()
p = doc.add_paragraph('The Question:')
run = p.runs[0]
run.bold = True
doc.add_paragraph('Do we need to overcomplicate this for production?')

doc.add_paragraph()
p = doc.add_paragraph('The Answer:')
run = p.runs[0]
run.bold = True
doc.add_paragraph('No.')

doc.add_paragraph()
doc.add_paragraph('For 200 ICT staff, where realistic concurrent usage is 2-5 people, a repurposed workstation is more than adequate. We don\'t need a £200k Kubernetes cluster. We don\'t need dual redundant servers. We don\'t need enterprise support contracts.')

doc.add_paragraph()
doc.add_paragraph('We need:')
doc.add_paragraph('• A server that works (repurposed workstation: £0)', style='List Bullet')
doc.add_paragraph('• Authentication so only ICT can access it (AD: already have)', style='List Bullet')
doc.add_paragraph('• Backups in case it breaks (automated: 1 hour setup)', style='List Bullet')
doc.add_paragraph('• Someone to keep an eye on it (David: 2 hours/month)', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Start simple. Prove value. Scale only if needed.')
run = p.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 94, 184)

doc.add_paragraph()
doc.add_paragraph('The beauty of this approach: if usage explodes and we actually need more power, upgrading is straightforward. The code works. The architecture is sound. We just move it to bigger hardware.')

doc.add_paragraph()
doc.add_paragraph('But let\'s be realistic: it\'s an HR policy bot. It\'ll get light, sporadic use. And that\'s fine. That\'s exactly what we designed for.')

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph('Recommended Next Steps:')
run = p.runs[0]
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph('1. Approval to proceed with pilot', style='List Bullet')
doc.add_paragraph('2. Identify repurposed hardware (or £11k budget for new)', style='List Bullet')
doc.add_paragraph('3. 2 weeks: David + senior engineer make it production-ready', style='List Bullet')
doc.add_paragraph('4. Launch to ICT, monitor for 3 months', style='List Bullet')
doc.add_paragraph('5. Review usage and decide: keep as-is, expand, or sunset', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph('Total Investment: £50/year + 2 weeks effort')
run = p.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 170, 0)

p = doc.add_paragraph('Total Savings vs Commercial: £35,850 over 3 years')
run = p.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 170, 0)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph('--- End of Deployment Plan ---')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.runs[0]
run.italic = True

doc.save('NHS_HR_RAG_Realistic_Deployment_Plan.docx')
print('✅ Realistic deployment plan created!')
print('📄 File: NHS_HR_RAG_Realistic_Deployment_Plan.docx')

