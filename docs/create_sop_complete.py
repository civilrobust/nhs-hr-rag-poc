from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NHS ICT HR Policy Assistant')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0, 94, 184)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Standard Operating Procedure')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 48, 135)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('System Startup & Shutdown Guide')
run.font.size = Pt(14)

# Metadata
doc.add_paragraph()
doc.add_paragraph('Version: 1.0')
doc.add_paragraph('Author: David Sewoke')
doc.add_paragraph('Date: March 2026')

doc.add_page_break()

# Section 1: Starting the System
doc.add_heading('Starting the System from PC Reboot', 1)

doc.add_heading('Prerequisites:', 2)
doc.add_paragraph('• PC is powered on and logged into Windows 11', style='List Bullet')
doc.add_paragraph('• User: civil', style='List Bullet')
doc.add_paragraph('• Hostname: WINDOWS-AAS0LVD', style='List Bullet')

doc.add_heading('Step 1: Open WSL Ubuntu Terminal', 2)
doc.add_paragraph('Method A (Recommended):')
doc.add_paragraph('1. Press Windows Key')
doc.add_paragraph('2. Type "Ubuntu"')
doc.add_paragraph('3. Click "Ubuntu" or "Ubuntu 24.04" application')

doc.add_paragraph('\nMethod B (Alternative):')
doc.add_paragraph('1. Press Windows Key + X')
doc.add_paragraph('2. Click "Windows Terminal" or "Terminal"')
doc.add_paragraph('3. Click dropdown arrow (v) at top')
doc.add_paragraph('4. Select "Ubuntu 24.04"')

doc.add_paragraph('\nYou should see:')
p = doc.add_paragraph('civil@WINDOWS-AAS0LVD:~$')
p.style = 'No Spacing'
run = p.runs[0]
run.font.name = 'Courier New'

doc.add_page_break()

doc.add_heading('Step 2: Start Backend Server', 2)

doc.add_paragraph('In Terminal 1, type these commands EXACTLY:')
doc.add_paragraph()

doc.add_paragraph('Command 1: Navigate to backend folder')
p = doc.add_paragraph('cd ~/nhs-hr-rag/backend')
run = p.runs[0]
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph('Command 2: Activate virtual environment')
p = doc.add_paragraph('source ../venv/bin/activate')
run = p.runs[0]
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph('You should see (venv) appear:')
p = doc.add_paragraph('(venv) civil@WINDOWS-AAS0LVD:~/nhs-hr-rag/backend$')
run = p.runs[0]
run.font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Command 3: Start backend server')
p = doc.add_paragraph('python3 api.py')
run = p.runs[0]
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph('✅ SUCCESS: Wait for this message:')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 170, 0)

p = doc.add_paragraph('INFO:     Uvicorn running on http://0.0.0.0:8000 (Press CTRL+C to quit)')
run = p.runs[0]
run.font.name = 'Courier New'

doc.add_paragraph()
p = doc.add_paragraph('⚠️ DO NOT CLOSE THIS TERMINAL! Leave it running.')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(255, 102, 0)

doc.add_page_break()

doc.add_heading('Step 3: Start Frontend Server', 2)

doc.add_paragraph('Open a NEW terminal:')
doc.add_paragraph('• Option A: Press Ctrl + Shift + T (new tab)')
doc.add_paragraph('• Option B: Open new Ubuntu window (repeat Step 1)')

doc.add_paragraph()
doc.add_paragraph('In Terminal 2, type:')
doc.add_paragraph()

doc.add_paragraph('Command 1: Navigate to frontend')
p = doc.add_paragraph('cd ~/nhs-hr-rag/frontend')
run = p.runs[0]
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph('Command 2: Start frontend')
p = doc.add_paragraph('npm run dev')
run = p.runs[0]
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph('✅ SUCCESS: Wait for:')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 170, 0)

p = doc.add_paragraph('VITE v4.4.0  ready in XXX ms')
run = p.runs[0]
run.font.name = 'Courier New'
p = doc.add_paragraph('  ➜  Local:   http://localhost:3000/')
run = p.runs[0]
run.font.name = 'Courier New'

doc.add_paragraph()
p = doc.add_paragraph('⚠️ DO NOT CLOSE THIS TERMINAL EITHER!')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(255, 102, 0)

doc.add_page_break()

doc.add_heading('Step 4: Open Browser', 2)

doc.add_paragraph('1. Open Chrome, Edge, or Firefox')
doc.add_paragraph('2. Type in address bar: http://localhost:3000')
doc.add_paragraph('3. Press Enter')

doc.add_paragraph()
p = doc.add_paragraph('✅ You should see the NHS ICT HR Policy Assistant!')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 170, 0)
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph('Features visible:')
doc.add_paragraph('• NHS logo and blue header', style='List Bullet')
doc.add_paragraph('• Statistics: 85 chunks, 10 policies', style='List Bullet')
doc.add_paragraph('• Example questions', style='List Bullet')
doc.add_paragraph('• Question input box at bottom', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('🎉 SYSTEM IS NOW RUNNING!')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.runs[0]
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 94, 184)

doc.add_page_break()

# Section 2: Stopping
doc.add_heading('Stopping the System', 1)

doc.add_paragraph('To properly shut down:')
doc.add_paragraph()

doc.add_paragraph('1. Go to Terminal 1 (Backend)')
doc.add_paragraph('2. Press Ctrl + C')
doc.add_paragraph('3. Wait for "Shutting down" message')
doc.add_paragraph()

doc.add_paragraph('4. Go to Terminal 2 (Frontend)')
doc.add_paragraph('5. Press Ctrl + C')
doc.add_paragraph()

doc.add_paragraph('6. Close both terminals')
doc.add_paragraph('7. Close browser tab')

doc.add_paragraph()
p = doc.add_paragraph('✅ System stopped safely')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 170, 0)

doc.add_page_break()

# Troubleshooting
doc.add_heading('Troubleshooting', 1)

doc.add_heading('Problem: Backend says "Address already in use"', 2)
doc.add_paragraph('Solution:')
p = doc.add_paragraph('killall python3')
run = p.runs[0]
run.font.name = 'Courier New'
doc.add_paragraph('Then start backend again')

doc.add_heading('Problem: Frontend says "Port 3000 in use"', 2)
doc.add_paragraph('Solution:')
p = doc.add_paragraph('killall node')
run = p.runs[0]
run.font.name = 'Courier New'
doc.add_paragraph('Then start frontend again')

doc.add_heading('Problem: Browser shows error', 2)
doc.add_paragraph('Check:')
doc.add_paragraph('1. Is backend running? (Terminal 1)')
doc.add_paragraph('2. Is frontend running? (Terminal 2)')
doc.add_paragraph('3. URL correct? http://localhost:3000')

doc.add_page_break()

# Quick Reference
doc.add_heading('Quick Reference Card', 1)

doc.add_heading('Complete Command Sequence', 2)

p = doc.add_paragraph('Terminal 1 (Backend):')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 94, 184)

p = doc.add_paragraph('cd ~/nhs-hr-rag/backend')
run = p.runs[0]
run.font.name = 'Courier New'
p = doc.add_paragraph('source ../venv/bin/activate')
run = p.runs[0]
run.font.name = 'Courier New'
p = doc.add_paragraph('python3 api.py')
run = p.runs[0]
run.font.name = 'Courier New'

doc.add_paragraph()
p = doc.add_paragraph('Terminal 2 (Frontend):')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 94, 184)

p = doc.add_paragraph('cd ~/nhs-hr-rag/frontend')
run = p.runs[0]
run.font.name = 'Courier New'
p = doc.add_paragraph('npm run dev')
run = p.runs[0]
run.font.name = 'Courier New'

doc.add_paragraph()
p = doc.add_paragraph('Browser:')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 94, 184)

p = doc.add_paragraph('http://localhost:3000')
run = p.runs[0]
run.font.name = 'Courier New'
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph('--- End of SOP ---')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.runs[0]
run.italic = True

doc.save('NHS_HR_RAG_SOP_COMPLETE.docx')
print('✅ Complete SOP created!')
print('📄 File: NHS_HR_RAG_SOP_COMPLETE.docx')

