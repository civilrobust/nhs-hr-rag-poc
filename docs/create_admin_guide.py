from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NHS ICT HR Policy Assistant')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0, 94, 184)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Administrator Guide')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 48, 135)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('Adding & Managing HR Policy Documents')
run.font.size = Pt(14)

# Metadata
doc.add_paragraph()
doc.add_paragraph('Version: 1.0')
doc.add_paragraph('Author: David Sewoke')
doc.add_paragraph('Date: March 2026')

doc.add_page_break()

# Overview
doc.add_heading('Overview', 1)
doc.add_paragraph('This guide explains how to add new HR policy documents to the system so the AI can answer questions about them.')

doc.add_paragraph()
doc.add_heading('How the AI "Learns" New Documents', 2)
doc.add_paragraph('The AI does not actually "learn" in the traditional sense. Instead, it uses a process called indexing:')
doc.add_paragraph()

doc.add_paragraph('1. Read: System reads all .txt files in the policies folder')
doc.add_paragraph('2. Split: Each document is split into chunks (paragraphs)')
doc.add_paragraph('3. Embed: Each chunk is converted to a 384-number vector (embedding)')
doc.add_paragraph('4. Index: Vectors are stored in ChromaDB for fast semantic search')
doc.add_paragraph('5. Search: When user asks question, system finds most relevant chunks')
doc.add_paragraph('6. Generate: Llama 3.1 8B reads chunks and generates answer')

doc.add_paragraph()
p = doc.add_paragraph('Important: This happens automatically every time the backend starts!')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 170, 0)

doc.add_page_break()

# Adding New Documents
doc.add_heading('Adding New HR Policy Documents', 1)

doc.add_heading('Step 1: Prepare Your Policy Document', 2)

doc.add_paragraph('Requirements:')
doc.add_paragraph('• Must be plain text format (.txt file)', style='List Bullet')
doc.add_paragraph('• UTF-8 encoding', style='List Bullet')
doc.add_paragraph('• Clear section breaks (double line breaks between paragraphs)', style='List Bullet')
doc.add_paragraph('• Structured format with headings', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Good Format Example:')
p = doc.add_paragraph('''NHS TRUST ICT DEPARTMENT
OVERTIME POLICY
Policy Reference: HR-ICT-011

1. POLICY STATEMENT
Clear paragraph explaining overtime policy.

2. ELIGIBILITY
Another clear paragraph about who is eligible.

3. RATES
Standard overtime: 1.5x hourly rate
Weekend overtime: 2x hourly rate''')
run = p.runs[0]
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_page_break()

doc.add_heading('Step 2: Add File to Policies Folder', 2)

doc.add_paragraph('Open WSL Ubuntu terminal and run:')
doc.add_paragraph()

p = doc.add_paragraph('cd ~/nhs-hr-rag/data/policies/')
run = p.runs[0]
run.font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Then copy your file here. You can:')
doc.add_paragraph()

doc.add_paragraph('Option A: Copy from Windows')
doc.add_paragraph('1. Open Windows File Explorer')
doc.add_paragraph('2. Navigate to: \\\\wsl$\\Ubuntu\\home\\civil\\nhs-hr-rag\\data\\policies\\')
doc.add_paragraph('3. Paste your .txt file here')

doc.add_paragraph()
doc.add_paragraph('Option B: Create directly in terminal')
p = doc.add_paragraph('nano new_policy.txt')
run = p.runs[0]
run.font.name = 'Courier New'
doc.add_paragraph('(Type your content, then Ctrl+O to save, Ctrl+X to exit)')

doc.add_paragraph()
doc.add_paragraph('Option C: Copy from another location')
p = doc.add_paragraph('cp /mnt/c/Users/YourName/Downloads/policy.txt .')
run = p.runs[0]
run.font.name = 'Courier New'

doc.add_page_break()

doc.add_heading('Step 3: Delete Old ChromaDB Index', 2)

p = doc.add_paragraph('⚠️ CRITICAL STEP: You must delete the old index!')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(255, 102, 0)

doc.add_paragraph()
doc.add_paragraph('Run this command:')
p = doc.add_paragraph('cd ~/nhs-hr-rag/data/')
run = p.runs[0]
run.font.name = 'Courier New'
p = doc.add_paragraph('rm -rf chroma_db/')
run = p.runs[0]
run.font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Why? ChromaDB stores indexed chunks. When you add new documents, the old index does not include them. Deleting forces a complete re-index.')

doc.add_paragraph()
p = doc.add_paragraph('✅ The chroma_db folder will be automatically recreated when backend starts')
run = p.runs[0]
run.font.color.rgb = RGBColor(0, 170, 0)

doc.add_page_break()

doc.add_heading('Step 4: Restart Backend to Re-Index', 2)

doc.add_paragraph('If backend is already running:')
doc.add_paragraph('1. Go to Terminal 1 (where backend is running)')
doc.add_paragraph('2. Press Ctrl + C to stop it')
doc.add_paragraph('3. Wait for shutdown')

doc.add_paragraph()
doc.add_paragraph('Then start it again:')
p = doc.add_paragraph('cd ~/nhs-hr-rag/backend')
run = p.runs[0]
run.font.name = 'Courier New'
p = doc.add_paragraph('source ../venv/bin/activate')
run = p.runs[0]
run.font.name = 'Courier New'
p = doc.add_paragraph('python3 api.py')
run = p.runs[0]
run.font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Watch the startup messages:')
p = doc.add_paragraph('''Loading policies from ../data/policies/
Loaded 95 chunks from 11 documents
INFO:     Uvicorn running on http://0.0.0.0:8000''')
run = p.runs[0]
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph('✅ If you see more chunks/documents than before, indexing worked!')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 170, 0)

doc.add_page_break()

doc.add_heading('Step 5: Verify in Browser', 2)

doc.add_paragraph('1. Open browser: http://localhost:3000')
doc.add_paragraph('2. Look at statistics on welcome screen')
doc.add_paragraph('3. Should show increased numbers:')
doc.add_paragraph('   • "95 Policy Chunks Indexed" (was 85)')
doc.add_paragraph('   • "11 HR Policies Loaded" (was 10)')

doc.add_paragraph()
doc.add_paragraph('4. Test by asking a question about the new policy')
doc.add_paragraph('   Example: "What is the overtime rate?"')

doc.add_paragraph()
p = doc.add_paragraph('✅ If AI answers correctly with citations from new policy, SUCCESS!')
run = p.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 170, 0)

doc.add_page_break()

# Complete Example
doc.add_heading('Complete Example: Adding Overtime Policy', 1)

doc.add_paragraph('Here is a complete walkthrough:')
doc.add_paragraph()

p = doc.add_paragraph('''# Step 1: Navigate to policies folder
cd ~/nhs-hr-rag/data/policies/

# Step 2: Create new policy file
cat > overtime_policy.txt << 'EOF'
NHS TRUST ICT DEPARTMENT
OVERTIME POLICY
Policy Reference: HR-ICT-011
Effective Date: January 2024

1. POLICY STATEMENT
Overtime may be authorized for eligible ICT staff when operational needs require work beyond normal hours.

2. ELIGIBILITY
All permanent ICT staff below Band 8 are eligible for overtime payments.

3. RATES
Standard overtime: 1.5x hourly rate
Weekend/Bank Holiday: 2x hourly rate

4. APPROVAL PROCESS
All overtime must be pre-approved by line manager via ESR system.

Contact: hr.ict@nhstrust.nhs.uk
EOF

# Step 3: Delete old index
cd ~/nhs-hr-rag/data/
rm -rf chroma_db/

# Step 4: Restart backend
cd ~/nhs-hr-rag/backend
source ../venv/bin/activate
python3 api.py

# Step 5: Test in browser at http://localhost:3000
# Ask: "What is the overtime rate for weekends?"
# Expected answer: "2x hourly rate" with citation to overtime_policy.txt''')
run = p.runs[0]
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_page_break()

# Removing Documents
doc.add_heading('Removing Policy Documents', 1)

doc.add_paragraph('To remove a policy from the system:')
doc.add_paragraph()

doc.add_paragraph('Step 1: Delete the file')
p = doc.add_paragraph('cd ~/nhs-hr-rag/data/policies/')
run = p.runs[0]
run.font.name = 'Courier New'
p = doc.add_paragraph('rm unwanted_policy.txt')
run = p.runs[0]
run.font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Step 2: Delete ChromaDB index')
p = doc.add_paragraph('cd ~/nhs-hr-rag/data/')
run = p.runs[0]
run.font.name = 'Courier New'
p = doc.add_paragraph('rm -rf chroma_db/')
run = p.runs[0]
run.font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Step 3: Restart backend (same as adding)')

doc.add_page_break()

# Troubleshooting
doc.add_heading('Troubleshooting', 1)

doc.add_heading('Problem: New document not showing in stats', 2)
doc.add_paragraph('Check:')
doc.add_paragraph('1. Is file in /data/policies/ folder?')
doc.add_paragraph('2. Is it a .txt file?')
doc.add_paragraph('3. Did you delete chroma_db folder?')
doc.add_paragraph('4. Did backend restart successfully?')

doc.add_heading('Problem: AI does not answer questions from new document', 2)
doc.add_paragraph('Check:')
doc.add_paragraph('1. Is document formatted with clear paragraphs?')
doc.add_paragraph('2. Are there double line breaks between sections?')
doc.add_paragraph('3. Try asking more specific questions')

doc.add_heading('Problem: Backend shows error during startup', 2)
doc.add_paragraph('Check:')
doc.add_paragraph('1. Is file valid UTF-8 text?')
doc.add_paragraph('2. No corrupted characters?')
doc.add_paragraph('3. Check backend terminal for error details')

doc.add_page_break()

# Technical Details
doc.add_heading('Technical Details', 1)

doc.add_heading('File Locations', 2)
doc.add_paragraph('• Policy documents: ~/nhs-hr-rag/data/policies/')
doc.add_paragraph('• ChromaDB index: ~/nhs-hr-rag/data/chroma_db/')
doc.add_paragraph('• Indexing code: ~/nhs-hr-rag/backend/rag_pipeline.py')

doc.add_heading('Indexing Process', 2)
doc.add_paragraph('1. Backend reads all .txt files in policies folder')
doc.add_paragraph('2. Each file split by double newlines (\\n\\n)')
doc.add_paragraph('3. Each chunk embedded using sentence-transformers model')
doc.add_paragraph('4. Embeddings stored in ChromaDB with metadata (filename, chunk index)')
doc.add_paragraph('5. Query time: user question embedded, top 3 similar chunks retrieved')
doc.add_paragraph('6. Retrieved chunks sent to Llama 3.1 8B for answer generation')

doc.add_heading('Chunk Size', 2)
doc.add_paragraph('• Optimal: 100-500 words per chunk')
doc.add_paragraph('• Current: Split by paragraph (\\n\\n)')
doc.add_paragraph('• Too small: Lacks context')
doc.add_paragraph('• Too large: Less precise matching')

doc.add_page_break()

# Quick Reference
doc.add_heading('Quick Reference Card', 1)

p = doc.add_paragraph('Add New Policy - Complete Command Sequence:')
run = p.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 94, 184)

doc.add_paragraph()

p = doc.add_paragraph('''cd ~/nhs-hr-rag/data/policies/
# Add your file here (copy from Windows or create)

cd ~/nhs-hr-rag/data/
rm -rf chroma_db/

cd ~/nhs-hr-rag/backend
# Stop backend if running (Ctrl+C)
source ../venv/bin/activate
python3 api.py

# Verify at http://localhost:3000''')
run = p.runs[0]
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph('--- End of Administrator Guide ---')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.runs[0]
run.italic = True

doc.save('NHS_HR_RAG_Administrator_Guide.docx')
print('✅ Administrator Guide created!')
print('📄 File: NHS_HR_RAG_Administrator_Guide.docx')

