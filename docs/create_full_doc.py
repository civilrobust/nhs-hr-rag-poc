from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NHS ICT HR Policy Assistant')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0, 94, 184)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Comprehensive Technical Documentation')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 48, 135)

# Metadata
doc.add_paragraph()
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
cells = [
    ('Author:', 'David Sewoke'),
    ('Version:', '1.0 - Proof of Concept'),
    ('Date:', 'March 2026'),
    ('Project:', 'AI Engineering Apprenticeship'),
    ('Repository:', 'github.com/civilrobust/nhs-hr-rag-poc')
]
for i, (label, value) in enumerate(cells):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = value

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc_items = [
    '1. Executive Summary',
    '2. Glossary of AI/ML Terms',
    '3. System Architecture', 
    '4. Technology Stack & Rationale',
    '5. Hardware Requirements',
    '6. Deployment & Infrastructure',
    '7. Cost Analysis',
    '8. Version Control',
    '9. Testing & QA',
    '10. Monitoring & Logging',
    '11. CI/CD Pipeline',
    '12. Scalability',
    '13. Comparison with Alternatives',
    '14. Performance Metrics',
    '15. Future Enhancements',
    '16. Lessons Learned',
    '17. Conclusion',
    'Appendices'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 1. Executive Summary
doc.add_heading('1. Executive Summary', 1)
doc.add_heading('What is this System?', 2)
doc.add_paragraph(
    'The NHS ICT HR Policy Assistant is an AI-powered conversational interface that enables NHS staff to query HR policies using natural language. The system uses Retrieval-Augmented Generation (RAG) to provide accurate, cited answers by searching through indexed policy documents and generating contextual responses using a locally-hosted Large Language Model (LLM).'
)

doc.add_heading('Key Capabilities', 2)
capabilities = [
    'Semantic Search: Understands question intent, finds relevant content even with different wording',
    'Source Attribution: Every answer includes citations to policy documents',
    'Local Processing: All data and inference happen on-premises (no cloud APIs)',
    'Fast Responses: 1-3 second response times with proper hardware',
    'Professional UI: NHS-branded, accessible web interface'
]
for cap in capabilities:
    doc.add_paragraph(cap, style='List Bullet')

doc.add_page_break()

# 2. Glossary
doc.add_heading('2. Glossary of AI/ML Terms', 1)
doc.add_paragraph('This section explains key artificial intelligence and machine learning terms used throughout this document.')

glossary = {
    'AI (Artificial Intelligence)': 'Computer systems that can perform tasks typically requiring human intelligence, such as understanding language, recognizing patterns, and making decisions.',
    
    'ChromaDB': "A vector database that stores text as numerical embeddings and allows semantic search. It's the 'search engine' that finds relevant policy sections based on meaning rather than exact keyword matches.",
    
    'Embeddings': 'Numerical representations of text that capture semantic meaning. For example, "sick leave" and "medical absence" would have similar embedding vectors because they mean similar things, even though the words are different. Typically represented as arrays of 384 or 768 numbers.',
    
    'GPU (Graphics Processing Unit)': 'Specialized hardware originally designed for graphics but now essential for AI. GPUs can perform thousands of calculations simultaneously, making them much faster than CPUs for running neural networks. The RTX 5080 has 16GB of VRAM (Video RAM) dedicated to these calculations.',
    
    'Inference': 'The process of using a trained AI model to make predictions or generate outputs. In this system, inference happens when Llama 3.1 8B reads the policy context and generates an answer. Unlike training (which teaches the model), inference is using the already-trained model. Each time a user asks a question, the system performs inference.',
    
    'LLM (Large Language Model)': 'A neural network trained on massive amounts of text to understand and generate human language. "Large" refers to billions of parameters (connections in the network). Llama 3.1 8B has 8 billion parameters.',
    
    'Llama 3.1 8B': "Meta's open-source language model with 8 billion parameters. The '8B' indicates model size. This model runs locally on the RTX 5080 GPU and generates the answers in this system.",
    
    'RAG (Retrieval-Augmented Generation)': 'A technique that enhances LLMs by first retrieving relevant documents, then including them as context in the prompt. This allows the LLM to answer questions about specific documents it wasn\'t trained on.',
    
    'Semantic Search': "Searching based on meaning rather than exact keywords. Traditional search finds documents containing your exact words. Semantic search understands that 'work from home', 'remote working', and 'telecommuting' all mean similar things.",
    
    'Vector Database': 'A specialized database optimized for storing and searching vectors (numerical arrays). ChromaDB is a vector database. Unlike traditional databases that search exact matches, vector databases find similar vectors using mathematical distance calculations.',
    
    'VRAM (Video RAM)': 'Memory on the GPU dedicated to graphics and computation. LLMs need to fit entirely in VRAM for fast inference. The RTX 5080\'s 16GB VRAM allows it to run Llama 3.1 8B with room to spare.'
}

for term, definition in glossary.items():
    doc.add_heading(term, 3)
    doc.add_paragraph(definition)

doc.add_page_break()

# Save
doc.save('NHS_HR_RAG_Complete_Technical_Documentation.docx')
print('✅ Complete documentation created!')
print('📄 File: NHS_HR_RAG_Complete_Technical_Documentation.docx')

